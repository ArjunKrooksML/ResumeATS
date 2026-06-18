from docx import Document


def find_header_footer_risks(document) -> list[str]:
    risks = []
    for section in document.sections:
        if any(p.text.strip() for p in section.header.paragraphs):
            risks.append("text found in header, which some ATS parsers ignore entirely")
        if any(p.text.strip() for p in section.footer.paragraphs):
            risks.append("text found in footer, which some ATS parsers ignore entirely")
    return risks


def find_table_risks(document) -> list[str]:
    if not document.tables:
        return []
    return [f"{len(document.tables)} table(s) detected, which ATS parsers often misread"]


def read_docx(file_path: str) -> str:
    document = Document(file_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    risks = find_table_risks(document) + find_header_footer_risks(document)
    if not risks:
        return text
    return f"{text}\n\nSTRUCTURE: {'; '.join(risks)}"
